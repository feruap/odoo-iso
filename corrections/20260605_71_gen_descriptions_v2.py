"""
Genera descripciones para cartuchos con formato:
"Cartucho de N ventanas para la prueba rápida de [desc] usado en muestras [muestra]"
Fuente: F-CC-007-001_full.docx
RE-EJECUTABLE.
"""
import re
from docx import Document

def count_ventanas(c_str):
    if not c_str or 'sin' in c_str.lower():
        return 1
    c_nums = re.findall(r'(\d+)\s*C', c_str)
    if c_nums:
        return sum(int(n) for n in c_nums)
    if 'C' in c_str.upper():
        return 1
    return 1

def clean(s):
    return s.strip().rstrip(',').rstrip('.').strip() if s else ''

doc = Document('/tmp/F-CC-007-001_full.docx')

cartucho_data = {}

# Tabla 0: MPCAR (col: 0=Clave, 2=Desc, 3=Muestra, 9=C, 10=T)
for row in doc.tables[0].rows[2:]:
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    if len(cells) < 11: continue
    key = cells[0].strip()
    if not key or key in cartucho_data: continue
    desc    = clean(cells[2])
    muestra = clean(cells[3])
    c_str   = cells[9]
    ventanas = count_ventanas(c_str)
    if desc and muestra:
        cartucho_data[key] = {'ventanas': ventanas, 'desc': desc, 'muestra': muestra}

# Tabla 1: MPCAC
for row in doc.tables[1].rows[2:]:
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    if len(cells) < 11: continue
    key = cells[0].strip()
    if not key or key in cartucho_data or not key.startswith('MPCAC'): continue
    desc    = clean(cells[2])
    muestra = clean(cells[3])
    c_str   = cells[9]
    ventanas = count_ventanas(c_str)
    if desc and muestra:
        cartucho_data[key] = {'ventanas': ventanas, 'desc': desc, 'muestra': muestra}

print(f"Datos del documento: {len(cartucho_data)} cartuchos")

templates = env['product.template'].search([('categ_id.name', '=', 'Cartucho')])
updated, skipped, not_found = [], [], []

for pt in templates:
    code = (pt.default_code or '').strip().upper()
    if not code:
        continue
    data = cartucho_data.get(code)
    if not data:
        not_found.append(code)
        continue

    v = data['ventanas']
    nueva = (
        f"Cartucho de {v} ventana{'s' if v != 1 else ''} "
        f"para la prueba rápida de {data['desc']} "
        f"usado en muestras {data['muestra']}"
    )

    if pt.description and pt.description.strip() == nueva:
        skipped.append(code)
        continue

    pt.write({'description': nueva})
    updated.append(f"  {code:<10} {v}V  {pt.name[:45]}")

env.cr.commit()

print(f"\n=== ACTUALIZADOS: {len(updated)} ===")
for l in updated[:15]: print(l)
if len(updated) > 15: print(f"  ...y {len(updated)-15} más")
print(f"\nSin datos en doc: {len(not_found)} | Ya correctos: {len(skipped)}")
print("Listo.")
