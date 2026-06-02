"""
Genera descripciones automáticas para cartuchos usando datos de F-CC-007-001.
Formato: "Cartucho de plástico de N ventanas para [detección] usado en muestras [tipo]"
"""
import re
from docx import Document

# ── 1. Parsear F-CC-007-001_full.docx ─────────────────────────────────────────

def count_c(c_str):
    if not c_str or 'sin' in c_str.lower() or 'no' in c_str.lower():
        return 0
    nums = re.findall(r'(\d+)\s*[CRcr]', c_str)
    if nums:
        return sum(int(n) for n in nums)
    if re.search(r'[CRcr]', c_str):
        return 1
    return 0

def count_t(t_str):
    # Quitar explicaciones en paréntesis: "3T (1 MYO/1 CKMB)" → "3T"
    t_clean = re.sub(r'\([^)]+\)', '', t_str).strip()
    # ¿Empieza con XTb? e.g. "5T" "3T" "2T"
    leading = re.match(r'(\d+)\s*T\b', t_clean)
    if leading:
        return int(leading.group(1))
    # Contar grupos separados por ";" y "," e.g. "1 IgG; 1 IgM" → 2
    groups = re.split(r'[;,]', t_clean)
    total = 0
    for g in groups:
        nums = re.findall(r'(\d+)', g)
        if nums:
            total += int(nums[0])
    return total if total > 0 else 1

def total_windows(c_str, t_str):
    return count_c(c_str) + count_t(t_str)

def clean_text(s):
    return s.strip().rstrip(',').rstrip('.').strip()

doc = Document('/tmp/F-CC-007-001_full.docx')

# TABLA 1: cartuchos individuales (MPCAR)
# Columnas: 0=Clave, 1=Nombre, 2=Descripción, 3=TipoMuestra, 9=C, 10=T
cartucho_data = {}  # code → {ventanas, muestra, descripcion}

t1 = doc.tables[0]
seen = set()
for row in t1.rows[2:]:
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    if len(cells) < 11:
        continue
    key = cells[0].strip()
    if not key or key in seen:
        continue
    seen.add(key)
    descripcion = clean_text(cells[2])
    muestra     = clean_text(cells[3])
    c_str       = cells[9]
    t_str       = cells[10]
    ventanas    = total_windows(c_str, t_str)
    cartucho_data[key] = {
        'ventanas': ventanas,
        'muestra':  muestra,
        'descripcion': descripcion,
    }

# TABLA 2: combos (MPCAC)
# Columnas: 0=Clave, 1=Nombre, 2=Descripción, 3=TipoMuestra, 10=C, 11=T(aprox)
t2 = doc.tables[1]
seen2 = set()
for row in t2.rows:
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    if len(cells) < 12:
        continue
    key = cells[0].strip()
    if not key or key in seen2 or not key.startswith('MPCAC'):
        continue
    seen2.add(key)
    descripcion = clean_text(cells[2])
    muestra     = clean_text(cells[3])
    c_str       = cells[9] if len(cells) > 9 else ''
    t_str       = cells[10] if len(cells) > 10 else ''
    ventanas    = total_windows(c_str, t_str)
    cartucho_data[key] = {
        'ventanas': ventanas,
        'muestra':  muestra,
        'descripcion': descripcion,
    }

# TABLA 3: generales (MPCAG) - "No aplica"
t3 = doc.tables[2]
seen3 = set()
for row in t3.rows:
    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
    if len(cells) < 12:
        continue
    key = cells[0].strip()
    if not key or key in seen3 or not key.startswith('MPCAG'):
        continue
    seen3.add(key)
    c_str    = cells[9] if len(cells) > 9 else ''
    t_str    = cells[10] if len(cells) > 10 else ''
    ventanas = total_windows(c_str, t_str)
    cartucho_data[key] = {
        'ventanas': ventanas,
        'muestra':  'No aplica',
        'descripcion': 'No aplica',
    }

print(f"Datos cargados de F-CC-007-001: {len(cartucho_data)} cartuchos")

# ── 2. Actualizar productos en Odoo ───────────────────────────────────────────

templates = env['product.template'].search([('categ_id.name', '=', 'Cartucho')])
print(f"Productos en Odoo: {len(templates)}")

updated = []
skipped_no_code = []
skipped_not_found = []
skipped_already = []

for pt in templates:
    code = (pt.default_code or '').strip().upper()
    if not code:
        skipped_no_code.append(pt.name)
        continue

    data = cartucho_data.get(code)
    if not data:
        skipped_not_found.append(f"{code} - {pt.name}")
        continue

    ventanas    = data['ventanas']
    muestra     = data['muestra']
    descripcion = data['descripcion']

    # Para MPCAG generales y "No aplica", no generar descripción automática
    if descripcion == 'No aplica' or muestra == 'No aplica':
        # Solo indicar ventanas si el nombre ya las menciona
        skipped_not_found.append(f"{code} - {pt.name} (general/no aplica)")
        continue

    nueva_desc = (
        f"Cartucho de plástico de {ventanas} ventana{'s' if ventanas != 1 else ''} "
        f"para {descripcion} "
        f"usado en muestras {muestra}"
    )

    # Verificar si ya tiene esa descripción
    if pt.description and pt.description.strip() == nueva_desc.strip():
        skipped_already.append(pt.name)
        continue

    pt.write({'description': nueva_desc})
    updated.append(f"  {code:<10} | {ventanas}V | {pt.name[:45]:<45} | {muestra[:35]}")

env.cr.commit()

# ── 3. Reporte ────────────────────────────────────────────────────────────────
print(f"\n=== ACTUALIZADOS: {len(updated)} ===")
for l in updated:
    print(l)

print(f"\n=== SIN DESCRIPCIÓN (no aplica/generales): {len(skipped_not_found)} ===")
for l in skipped_not_found:
    print(f"  {l}")

print(f"\n=== Sin código interno: {len(skipped_no_code)} ===")
for l in skipped_no_code:
    print(f"  {l}")

print(f"\n=== Ya tenían descripción correcta: {len(skipped_already)} ===")
print("Listo.")
